# -*- coding: utf-8  -*-
# -*- author: jokker -*-

from Report.WordUtil import WordUtil
import docx

# doc_path = r'C:\Users\Administrator\Desktop\wordUtil_test\依赖注入.doc'
# a = docx.Document(doc_path)
# a = WordUtilNew(r'C:\Users\Administrator\Desktop\wordUtil_test\依赖注入.doc')


a = WordUtil()

font_style_new = a.get_font_style(font_name=u'宋体', font_size=14, font_color=(0, 0, 0))
font_style_2 = a.get_font_style(font_name=u'宋体', font_size=14, font_color=(0, 0, 0), font_is_bold=True)
font_style_3 = a.get_font_style(font_name=u'黑体', font_size=14, font_color=(0, 100, 255), font_is_italic=True)
font_style_4 = a.get_font_style(font_name=u'微软雅黑', font_size=14, font_color=(255, 0, 0))
font_style_5 = a.get_font_style(font_name=u'微软雅黑', font_size=10, font_color=(30, 30, 255))

# 测试添加文字段落
a.add_paragraph([{'text': '    要是想在段落前面空两行的话，可以直接在段落的 text 字符串前面加四个空格，'
                          '四个空格相当于两个中文字符', 'font_style': font_style_new}])

a.add_paragraph([{'text': '', 'font_style': font_style_new}])  # 插入一个空的段落当做回车键使用

a.add_paragraph([{'text': '    可以对一个段落中的字符串使用多个样式。', 'font_style': font_style_new},
                 {'text': '第一个样式，宋体，14号，黑色，加粗；', 'font_style': font_style_2},
                 {'text': '第二个样式，黑体，14号，蓝色，不加粗，斜体；', 'font_style': font_style_3},
                 {'text': '第二个样式，微软雅黑，14号，红色，不加粗，不斜体。', 'font_style': font_style_4}, ])

a.add_paragraph([{'text': '', 'font_style': font_style_new}])  # 插入一个空的段落当做回车键使用

p1 = a.add_paragraph([{'text': '可以调整段落样式', 'font_style': font_style_new},
                      {'text': '居左', 'font_style': font_style_4},
                      {'text': '显示', 'font_style': font_style_new}])

p2 = a.add_paragraph([{'text': '可以调整段落样式', 'font_style': font_style_new},
                      {'text': '居中', 'font_style': font_style_4},
                      {'text': '显示', 'font_style': font_style_new}])

p3 = a.add_paragraph([{'text': '可以调整段落样式', 'font_style': font_style_new},
                      {'text': '居右', 'font_style': font_style_4},
                      {'text': '显示', 'font_style': font_style_new}])

a.set_paragraph_style(p1, 'LEFT')
a.set_paragraph_style(p2, 'CENTER')
a.set_paragraph_style(p3, 'RIGHT')

# 测试添加表格
a.add_paragraph([{'text': '', 'font_style': font_style_new}])  # 插入一个空的段落当做回车键使用

cell_style_new = a.get_cell_style(font_style_new, 'LEFT', True)

table_info = [[{'text': 'num_1', 'cell_style': cell_style_new},
               {'text': 'num_2', 'cell_style': cell_style_new},
               {'text': 'num_3', 'cell_style': cell_style_new},
               {'text': 'num_4', 'cell_style': cell_style_new}],

              [{'text': 'value_1', 'cell_style': cell_style_new},
               {'text': 'value_2', 'cell_style': cell_style_new},
               {'text': 'value_3', 'cell_style': cell_style_new},
               {'text': 'value_4', 'cell_style': cell_style_new}],
              ]

merge_info = [{'merge_value': u'支持单元格合并', 'merge_range': (0, 0, 1, 1),
               'cell_style': a.get_cell_style(font_style_5, 'LEFT', True)},
              {'merge_value': u'支持单元格合并', 'merge_range': (0, 3, 1, 3),
               'cell_style': a.get_cell_style(font_style_5, 'RIGHT', False)}]

a.add_table(table_info, merge_info)

# 添加换页符
a.add_paragraph([{'text': '', 'font_style': font_style_new}])
a.add_paragraph([{'text': '可以添加换页符，这是换页符之前', 'font_style': font_style_new}])
a.add_page_break()
a.add_paragraph([{'text': '这是换页符之后', 'font_style': font_style_new}])

# 添加图片
P0 = a.add_paragraph([{'text': '下面增加两个自定义大小的图片', 'font_style': font_style_new}])
a.set_paragraph_style(P0, alignment='CENTER')
a.add_picture(r'.\word\12.jpg', pic_width=8, alignment='CENTER')
p1 = a.add_paragraph([{'text': '空一行', 'font_style': font_style_new}])
a.set_paragraph_style(p1, alignment='CENTER')
a.add_picture(r'.\word\12.jpg', pic_width=10, alignment='CENTER')

# 替换文字
a.replace_word({'佐樱': 'other_woman'})

# 表格内容进行修改
font_style_new = a.get_font_style(font_name=u'宋体', font_size=25, font_color=(255, 0, 0), font_is_bold=True)
cell_style_new = WordUtil.get_cell_style(font_style_new)
a.add_lines_in_assign_table(0, [[{'text': None, 'cell_style': cell_style_new}, {'text': None, 'cell_style': None},
                                 {'text': None, 'cell_style': None}, {'text': None, 'cell_style': None}],
                                [{'text': None, 'cell_style': None}, {'text': None, 'cell_style': cell_style_new},
                                 {'text': None, 'cell_style': cell_style_new}, {'text': None, 'cell_style': cell_style_new}],
                                [{'text': '支持对指定表格修改', 'cell_style': cell_style_new},
                                 {'text': '增加行列', 'cell_style': cell_style_new},
                                 {'text': '改变样式', 'cell_style': cell_style_new},
                                 {'text': '或者忽略', 'cell_style': cell_style_new}]])

# 保存 data 文档
a.save(r'C:\Users\Administrator\Desktop\del.doc')
