# -*- coding: utf-8  -*-
# -*- author: jokker -*-

import docx
from docxtpl import DocxTemplate
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt


# 参考 ： https://blog.csdn.net/sinat_30711195/article/details/80725435
# 参考 ： https://blog.csdn.net/shiyuzuxiaqianli/article/details/97043903

# todo 图片的替换暂时使用 WordUtilOld 来实现，也就是模板实现的时候 WordUtil，WordUtilOld 一起使用

# todo 增加表格左右对齐的功能，找到地方了，但是电脑太慢以后再写

# todo 表格样式修改部分还需要进行完善

class WordUtilOld(object):

    def __init__(self, doc_template, replace_dict, save_path):
        # 模板
        self.doc_template = doc_template
        # 替换字典
        self.replace_dict = replace_dict
        # 保存的字典
        self.save_path = save_path
        #
        self.__tpl = DocxTemplate(self.doc_template)

    @staticmethod
    def __replace_text(tpl, old_text, new_text):
        for p in tpl.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in inline:
                    if old_text in i.text:
                        text = i.text.replace(old_text, new_text)
                        i.text = text

    def __replace_pics(self):
        """替换图片"""
        if 'pics' in self.replace_dict:
            replace_pic = self.replace_dict['pics']
            for each in replace_pic:
                # 替换
                print('replace pics | {0} | {1}'.format(each, replace_pic[each]))
                self.__tpl.replace_media(each, replace_pic[each])

    def __replace_words(self):
        """替换文字"""
        # render 要两个括号括住
        if 'words' in self.replace_dict:
            self.__tpl.render(self.replace_dict['words'])

    def __replace_tables(self):
        """替换表格"""
        if 'tables' not in self.replace_dict:
            return

        replace_table = self.replace_dict['tables']
        # 遍历每一个表格
        for each_table in replace_table:
            # ----------------------------------------------------------------------------------------------------------
            # 每次增加一行的模式
            if each_table['mode'] == 'add_row':
                table = self.__tpl.tables[each_table['num']]
                for line_index, each_line in enumerate(each_table['value']):
                    # 表格增加一行
                    row_cells = table.add_row().cells
                    for element_index, each_element in enumerate(each_line):
                        # 将 float 和 int 转为 str，而不是全部转为 str 防止出现 unicode 的情况
                        if isinstance(each_element, int) or isinstance(each_element, float):
                            row_cells[element_index].text = str(each_element)
                        else:
                            row_cells[element_index].text = each_element
            # ----------------------------------------------------------------------------------------------------------
            # 表格不变换，替换其中需要替换的内容
            # FIXME 替换有问题，需要解决
            elif each_table['mode'] == 'replace':
                # FIXME 如何才能只替换当前表格中的元素呢？如何过滤指定表格中的内容和其他部分的内容
                if 'words' in self.replace_dict:
                    self.replace_dict['words'].update(each_table['value'])
                else:
                    self.replace_dict['words'] = each_table['value']
                # print('table {0} : replace mode'.format(each_table['num']))
            # ----------------------------------------------------------------------------------------------------------
            else:
                print('replace table error')

    def __save_docx(self):
        """保存到文件"""
        self.__tpl.save(self.save_path)
        # print('new docx save success : {0}'.format(self.save_path))

    def do_process(self):
        # 替换图片
        self.__replace_pics()
        # 替换表格
        self.__replace_tables()
        # 替换文字
        self.__replace_words()
        # 保存
        self.__save_docx()


class WordUtil(object):
    """新的 data 操作方法"""

    # todo 调整自动生成表格的上下边线是否显示
    # todo 调整段落之间的行间距
    # todo 设置单元格长宽
    # todo 增加读 data 的功能
    # todo 增加插入新内容的功能
    # todo 插入一个 段落（图片、文字）
    # todo 当样式为 None 的时候使用默认的样式
    # todo 添加读取 data 中 table 的功能(×)

    def __init__(self, doc_path=None):
        if doc_path is None:
            self.__document = docx.Document()
        else:
            self.__document = docx.Document(doc_path)
    # ------------------------------------- 段落文字设置 -----------------------------------------------------
    @staticmethod
    def __set_run_style(run_obj, font_style):
        """设置 run 对象的格式，可以自定义段落中的某一些文字的样式"""
        # 字体的样式
        font_color = font_style['font_color']
        font_size = font_style['font_size']
        font_name = font_style['font_name']
        font_is_bold = font_style['font_is_bold']
        font_is_italic = font_style['font_is_italic']
        # 设置字体的样式
        run_obj.font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])  # 字体 rgb 颜色设置
        run_obj.font.size = Pt(float(font_size))  # 字体大小
        run_obj.bold = font_is_bold  # 是否加粗
        run_obj.font.italic = font_is_italic  # 是否斜体
        run_obj.font.name = font_name
        run_obj._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # fixme 弄不明白改一下字体的颜色，为什么用上面两句话

    @staticmethod
    def set_paragraph_style(p, alignment='center', ):
        """设置段落的格式"""
        if alignment.upper() == 'CENTER':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment.upper() == 'LEFT':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment.upper() == 'RIGHT':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    @staticmethod
    def get_font_style(font_color=(0, 0, 0), font_size=20, font_is_bold=False, font_is_italic=False,
                       font_name=u'宋体'):
        """获取字体的样式"""
        font_style = {'font_color': font_color, 'font_size': font_size, 'font_is_bold': font_is_bold,
                      'font_is_italic': font_is_italic, 'font_name': font_name}
        return font_style

    def add_paragraph(self, paragraph_info):
        # 传入的应该是 一个个 run 拼接而成的  paragraph ， 所以需要指定每个 run的 格式
        # DS --> [{'text':'text_value', 'font_style': font_style}] --> 内容 和 内容对应的样式
        """增加一个段落，可以设置段落的格式"""
        p = self.__document.add_paragraph()
        for each_run_info in paragraph_info:
            each_text = each_run_info['text']
            each_font_style = each_run_info['font_style']
            each_run = p.add_run(each_text)
            self.__set_run_style(each_run, each_font_style)
        return p  # 返回 p 对象，用于更多的操作
    # -------------------------------------- 表格设置 --------------------------------------------------------
    @staticmethod
    def set_table_style(table_obj, alignment='center', rgb_color=(0, 0, 0)):
        """获取表格样式数据"""
        table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中

        if alignment.upper() == 'CENTER':
            table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif alignment.upper() == 'LEFT':
            table_obj.alignment = WD_TABLE_ALIGNMENT.LEFT
        elif alignment.upper() == 'RIGHT':
            table_obj.alignment = WD_TABLE_ALIGNMENT.RIGHT
        # 设置表格的颜色
        table_obj.style.font.color.rgb = RGBColor(rgb_color[0], rgb_color[1], rgb_color[2])  # 表格整体改变颜色

    @staticmethod
    def set_cell_text_and_style(cell, cell_text, cell_style):
        """设置单元格的样式和内容"""
        # 设置单元格垂直居中
        if 'is_center_vertical' in cell_style:
            if cell_style['is_center_vertical']:
                cell.vertical_alignment = True  # 默认所有的单元格都是垂直居中的

        pa = cell.paragraphs[0]  # 表中的第一个段落
        # 设置单元格水平居中
        if 'is_center_horizontal' in cell_style:

            print(cell_style['is_center_horizontal'])
            
            # todo 这个地方改一下，这个参数名很明显是有问题的

            if str(cell_style['is_center_horizontal']).upper() == 'LEFT':
                pa.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if str(cell_style['is_center_horizontal']).upper() == 'RIGHT':
                pa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            if cell_style['is_center_horizontal'] is True:
                pa.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中

        # 可以直接设置每一个 cell 中的 text 和对应的 font_style
        if cell_text is not None:
            run = pa.add_run(cell_text)
            WordUtil.__set_run_style(run, cell_style)

    @staticmethod
    def set_table_merge_info(table, merge_info):
        """设置单元格合并信息"""
        for each_merge_info in merge_info:
            merge_value = each_merge_info['merge_value']
            merge_range = each_merge_info['merge_range']
            cell_style = each_merge_info['cell_style']
            # 合并单元格
            table.cell(merge_range[0], merge_range[1]).merge(table.cell(merge_range[2], merge_range[3]))
            #
            each_cell = table.cell(merge_range[0], merge_range[1])
            # 对单元格的格式进行清空
            each_cell.text = ''
            pa = each_cell.paragraphs[0]  # 表中的第一个段落

            if 'is_center_vertical' in cell_style:
                if cell_style['is_center_vertical']:
                    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中

            # 设置单元格垂直居中
            if 'is_center_horizontal' in cell_style:

                # print(cell_style['is_center_horizontal'])
                # print(cell_style['is_center_vertical'])

                # if (str(cell_style['is_center_horizontal']).upper()) == 'LEFT'.upper():
                # if True:
                #     pa.alignment = WD_ALIGN_PARAGRAPH.LEFT
                #
                # if str(cell_style['is_center_horizontal']).upper() == 'RIGHT':
                #     pa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                #
                # if cell_style['is_center_horizontal'] is True:
                #     pa.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中

                pa.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 可以直接设置每一个 cell 中的 text 和对应的 font_style
            run = pa.add_run(merge_value)
            WordUtil.__set_run_style(run, cell_style)

    @staticmethod
    def get_cell_style(font_style=None, is_center_horizontal=True, is_center_vertical=True):
        """获取表格单元的格式"""
        cell_style = {'font_color': (0, 0, 0), 'font_size': 20, 'font_is_bold': False,
                      'font_is_italic': False, 'font_name': u'宋体', 'is_center_horizontal': True,
                      'is_center_vertical': True}

        if font_style is not None:
            for each_style in ['font_color', 'font_size', 'font_is_bold', 'font_is_italic', 'font_name']:
                if each_style in font_style:
                    cell_style[each_style] = font_style[each_style]

        if is_center_horizontal:
            cell_style['is_center_horizontal'] = is_center_horizontal
        else:
            cell_style['is_center_horizontal'] = False

        if is_center_vertical:
            cell_style['is_center_vertical'] = is_center_vertical
        else:
            cell_style['is_center_vertical'] = False

        return cell_style

    def add_table(self, table_info, merge_info=None, style='Table Grid'):
        """增加一个表格，返回 table，通过操作 table 来改变表格的内容和样式
        table_info：DS --> [[{'text': 'test_text', 'cell_style': cell_style}], []]
        merge_info：DS --> [{merge_value, merge_range, cell_style},]"""
        row_num = len(table_info)  # 查看表格的行数
        if row_num == 0:
            return
        else:
            column_num = len(table_info[0])  # 列的个数

        table = self.__document.add_table(rows=row_num, cols=column_num, style=style)

        # 设置每一个单元格
        for i in range(row_num):
            for j in range(column_num):
                cell = table.cell(i, j)
                cell_text = table_info[i][j]['text']
                cell_style = table_info[i][j]['cell_style']
                WordUtil.set_cell_text_and_style(cell, cell_text, cell_style)

        WordUtil.set_table_merge_info(table, merge_info)  # 设置合并
        return table
    # --------------------------------------------- 替换功能设置 -------------------------------------------------------

    def replace_word(self, replace_dict):
        """替换文字"""
        # todo 考虑替换文字的时候是否需要指定样式
        for para in self.__document.paragraphs:
            for i in range(len(para.runs)):
                for key, value in replace_dict.items():
                    if key in para.runs[i].text:
                        print(key + "->" + value)
                        para.runs[i].text = para.runs[i].text.replace(key, value)

    def add_lines_in_assign_table(self, table_index, lines_info, merge_info=None):
        """对于指定的表格增加行"""
        table = self.__document.tables[table_index]
        # 如果行列超过要求那么就新建需要的行列
        table_line_num = len(table.rows)
        table_column_num = len(table.columns)

        for line_index, each_line in enumerate(lines_info):
            if line_index > table_line_num - 1:
                table.add_row()
            for column_index, each_cell_info in enumerate(each_line):
                if column_index > table_column_num - 1:
                    table.add_column()

                cell = table.cell(line_index, column_index)
                cell_text = each_cell_info['text']  #
                cell_style = WordUtil.get_cell_style()  # 设置表格默认样式

                if 'cell_style' in each_cell_info:
                    if each_cell_info['cell_style'] is not None:
                        cell_style = each_cell_info['cell_style']
                WordUtil.set_cell_text_and_style(cell, cell_text, cell_style)

        if merge_info is not None:
            WordUtil.set_table_merge_info(table, merge_info)  # 设置表格的合并信息
    # --------------------------------------------------------------------------------------------------------

    def get_assign_paragraph_info(self, assign_index):
        """获取指定段落的信息"""
        paragraph = self.__document.paragraphs[assign_index]
        return {'text': paragraph.text, 'style': paragraph.style}

    def insert_paragraph(self):
        """插入一个段落"""

    def add_page_break(self):
        """增加分页符"""
        self.__document.add_page_break()

    def add_heading(self, title_str, title_level=2):
        """增加标题"""
        self.__document.add_heading(title_str, level=title_level)

    def add_picture(self, pic_path, pic_width=10, alignment='center'):
        """插入图片"""
        self.__document.add_picture(pic_path, width=docx.shared.Cm(pic_width))  # 设置宽度，单位是厘米
        # 设置图片居中显示
        last_paragraph = self.__document.paragraphs[-1]  # 图片可以看做是最后一个段落

        if alignment.upper() == 'CENTER':
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment.upper() == 'LEFT':
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment.upper() == 'RIGHT':
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def save(self, save_path):
        """保存的路径"""
        self.__document.save(save_path)

    def to_do(self):
        """设置"""

        # todo 设置节的四边距离
        distance = Inches(0.3)
        sec = self.__document.sections[0]  # sections对应文档中的“节”
        sec.left_margin = distance  # 以下依次设置左、右、上、下页面边距
        sec.right_margin = distance
        sec.top_margin = distance
        sec.bottom_margin = distance

        # todo 设置页面的宽度和高度
        sec.page_width = Inches(12)  # 设置页面宽度
        sec.page_height = Inches(20)  # 设置页面高度

        # todo 在表格中插入图片
        # paragraph = cell.paragraphs[0]
        # run = paragraph.add_run()
        # run.add_picture('image.png')

        #

        # --------------------------------------------------





